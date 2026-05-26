#!/usr/bin/env python3
# pyright: reportMissingImports=false, reportUnknownVariableType=false, reportUnknownMemberType=false, reportUnknownParameterType=false, reportUnknownArgumentType=false, reportAny=false, reportDeprecated=false, reportUnusedCallResult=false
"""
Replace Mermaid blocks in DOCX with rendered PNG images.

Usage:
  .venv/bin/python docx_mermaid_to_png.py input.docx -o output.docx
"""

from __future__ import annotations

import argparse
import hashlib
import subprocess
import tempfile
from collections.abc import Iterable
from pathlib import Path

from docx import Document
from docx.document import Document as _Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


MERMAID_START_KEYWORDS = (
    "graph ",
    "flowchart ",
    "sequenceDiagram",
    "classDiagram",
    "stateDiagram",
    "stateDiagram-v2",
    "erDiagram",
    "journey",
    "gantt",
    "pie",
    "mindmap",
    "timeline",
    "gitGraph",
    "quadrantChart",
    "requirementDiagram",
    "sankey-beta",
    "c4context",
    "c4container",
    "c4component",
    "c4dynamic",
    "c4deployment",
)


def iter_block_items(parent: _Document | _Cell) -> Iterable[Paragraph | Table]:
    parent_elm = parent.element.body if isinstance(parent, _Document) else parent._tc
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def iter_paragraphs(parent: _Document | _Cell) -> Iterable[Paragraph]:
    for block in iter_block_items(parent):
        if isinstance(block, Paragraph):
            yield block
        else:
            for row in block.rows:
                for cell in row.cells:
                    yield from iter_paragraphs(cell)


def normalize_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()


def is_mermaid_block(text: str) -> bool:
    normalized = normalize_text(text)
    if not normalized:
        return False

    lines = [line.rstrip() for line in normalized.split("\n")]
    first_non_empty = next((line.strip() for line in lines if line.strip()), "")
    if not first_non_empty:
        return False

    first_lower = first_non_empty.lower()
    if not any(first_lower.startswith(prefix.lower()) for prefix in MERMAID_START_KEYWORDS):
        return False

    # Exclude false positives: Mermaid block should usually have more content than one token.
    return len(lines) > 1 or "-->" in normalized or ":::" in normalized or "{" in normalized


def render_mermaid_png(code: str, output_png: Path, use_global_mmdc: bool) -> None:
    with tempfile.NamedTemporaryFile("w", suffix=".mmd", encoding="utf-8", delete=False) as tmp:
        tmp.write(code)
        tmp_mmd_path = Path(tmp.name)

    try:
        if use_global_mmdc:
            cmd = ["mmdc", "-i", str(tmp_mmd_path), "-o", str(output_png)]
        else:
            cmd = [
                "npx",
                "-y",
                "@mermaid-js/mermaid-cli",
                "-i",
                str(tmp_mmd_path),
                "-o",
                str(output_png),
            ]

        result = subprocess.run(cmd, check=False, capture_output=True, text=True)
        if result.returncode != 0:
            msg = (
                "Mermaid render failed:\n"
                + f"cmd: {' '.join(cmd)}\n"
                + f"stdout: {result.stdout}\n"
                + f"stderr: {result.stderr}"
            )
            raise RuntimeError(msg)

        if not output_png.exists() or output_png.stat().st_size == 0:
            raise RuntimeError(f"Mermaid CLI did not produce PNG: {output_png}")
    finally:
        tmp_mmd_path.unlink(missing_ok=True)


def clear_paragraph(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def replace_mermaid_blocks(
    doc: _Document,
    temp_png_dir: Path,
    image_width_in: float,
    use_global_mmdc: bool,
) -> int:
    rendered_cache: dict[str, Path] = {}
    replaced_count = 0

    for paragraph in iter_paragraphs(doc):
        content = normalize_text(paragraph.text)
        if not is_mermaid_block(content):
            continue

        code_hash = hashlib.sha1(content.encode("utf-8")).hexdigest()[:16]
        png_path = rendered_cache.get(code_hash)
        if png_path is None:
            png_path = temp_png_dir / f"mermaid-{code_hash}.png"
            render_mermaid_png(content, png_path, use_global_mmdc=use_global_mmdc)
            rendered_cache[code_hash] = png_path

        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(png_path), width=Inches(image_width_in))
        replaced_count += 1

    return replaced_count


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Replace Mermaid blocks in DOCX with PNG images rendered by Mermaid CLI."
    )
    parser.add_argument("input", type=Path, help="Input .docx file path")
    parser.add_argument("-o", "--output", type=Path, help="Output .docx file path")
    parser.add_argument(
        "--image-width-in",
        type=float,
        default=6.4,
        help="Inserted image width in inches (default: 6.4)",
    )
    parser.add_argument(
        "--use-global-mmdc",
        action="store_true",
        help="Use globally installed mmdc command instead of npx mermaid-cli.",
    )
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()

    input_docx = args.input
    if not input_docx.exists():
        parser.error(f"Input file does not exist: {input_docx}")
    if input_docx.suffix.lower() != ".docx":
        parser.error("Input file must be a .docx")

    output_docx = args.output or input_docx.with_name(f"{input_docx.stem}.png-replaced.docx")

    temp_png_dir = Path(tempfile.mkdtemp(prefix="docx-mermaid-png-"))
    doc = Document(str(input_docx))

    replaced = replace_mermaid_blocks(
        doc=doc,
        temp_png_dir=temp_png_dir,
        image_width_in=args.image_width_in,
        use_global_mmdc=args.use_global_mmdc,
    )

    doc.save(str(output_docx))

    print(f"Input: {input_docx}")
    print(f"Output: {output_docx}")
    print(f"Mermaid blocks replaced: {replaced}")
    print(f"Rendered PNG cache dir: {temp_png_dir}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
